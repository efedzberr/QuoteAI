import os
import re
import time
import difflib
import json
import unicodedata
import math
import io
import base64
import threading
from datetime import datetime, timezone
from collections import Counter
from concurrent.futures import ThreadPoolExecutor
from flask import Flask, request, Response
from flask_cors import CORS

# ===========================================================================
# SERVICIO ÚNICO EN RAILWAY
# ---------------------------------------------------------------------------
# Este archivo consolida TODO el backend del flujo de cotizaciones:
#
#   /extract          -> NUEVO. Recibe el archivo (multipart 'file') y devuelve
#                        el JSON canónico de productos. Reemplaza a la Edge
#                        Function de Supabase y a la decodificación que vivía en
#                        el navegador (SheetJS / pdf.js / mammoth).
#
#                          • xlsx/xls/csv  -> decodificación determinística en
#                                             Python (sin LLM, rápido).
#                          • png/jpg/...   -> Claude visión directo.
#                          • pdf con texto -> extracción local + LLM.
#                          • pdf escaneado -> Docling (servicio HTTP) + fallback
#                                             a Claude documento.
#                          • docx          -> texto local + LLM.
#
#   /match            -> SIN CAMBIOS. Matching contra catálogo + validación LLM.
#   /catalog/refresh  -> SIN CAMBIOS.
#   /health           -> SIN CAMBIOS (se le añade el estado del extractor).
#
# Con CORS habilitado, Bolt llama a /extract y /match DIRECTAMENTE. Ya no hace
# falta el escenario de Make (era solo un relevo) ni la Edge Function.
# ===========================================================================

app = Flask(__name__)

# --- CORS ------------------------------------------------------------------
# Pon en CORS_ORIGINS (env) el dominio de Bolt separado por comas, p. ej.
#   CORS_ORIGINS=https://miapp.bolt.host,https://otra.com
# Mientras pruebas puedes dejar "*". No usamos cookies, así que "*" es seguro.
_cors_env = os.environ.get("CORS_ORIGINS", "*").strip()
_cors_origins = "*" if _cors_env == "*" else [o.strip() for o in _cors_env.split(",") if o.strip()]
CORS(app, resources={r"/*": {"origins": _cors_origins}})

SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY")
# Para ESCRIBIR avance y job_items en segundo plano, Railway necesita la
# service_role key (omite RLS). Si no se define, cae a SUPABASE_KEY.
SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY") or SUPABASE_KEY
# Tamaño de lote del trabajo asíncrono de matching (cada lote = un avance).
JOB_CHUNK = int(os.environ.get("JOB_CHUNK", "30"))

# ---------------------------------------------------------------------------
# Sinónimos de unidades técnicas
# ---------------------------------------------------------------------------
UNIT_SYNONYMS = {
    'KW': ['KW', 'KILOWATT', 'KILOWATTS'],
    'HP': ['HP', 'HORSEPOWER', 'CABALLOS'],
    'A':  ['A', 'AMP', 'AMPS', 'AMPERE', 'AMPERES', 'AMPER'],
    'V':  ['V', 'VOLT', 'VOLTS', 'VOLTIOS'],
    'HZ': ['HZ', 'HERTZ', 'HERCIOS'],
    'MM': ['MM', 'MILIMETRO', 'MILIMETROS'],
    'PZ': ['PZ', 'PZA', 'PIEZA', 'PIEZAS'],
    '3P': ['3P', 'TRIPOLAR', 'TRIFASICO', 'TRI', '3 POLOS', '3POLOS'],
    '2P': ['2P', 'BIPOLAR', 'BIFASICO', '2 POLOS', '2POLOS'],
}
SYNONYM_MAP = {v: k for k, variants in UNIT_SYNONYMS.items() for v in variants}

# ---------------------------------------------------------------------------
# Palabras vacías — se ignoran en la búsqueda por palabras clave
# ---------------------------------------------------------------------------
STOPWORDS = {
    'DE', 'DEL', 'LA', 'LAS', 'EL', 'LOS', 'UN', 'UNA', 'UNOS', 'UNAS',
    'Y', 'O', 'E', 'U', 'A', 'PARA', 'POR', 'CON', 'SIN', 'EN', 'AL',
    'SE', 'NO', 'SI', 'QUE', 'SU', 'MAS', 'ES', 'SON', 'THE', 'OF',
    'AND', 'OR', 'FOR', 'TO', 'IN', 'AT', 'BY',
}

# ---------------------------------------------------------------------------
# Posibles nombres de columna donde puede venir el código del producto
# ---------------------------------------------------------------------------
CODE_COLUMN_NAMES = [
    'codigoart', 'codigo', 'code', 'sku', 'clave',
    'cve', 'cve_art', 'cveart', 'no_parte', 'noparte',
    'numparte', 'partnumber', 'part_number', 'item_code',
    'itemcode', 'product_code', 'productcode', 'ref', 'referencia'
]

# Nombres de columna para cantidad y unidad (decodificación estructurada)
QTY_COLUMN_NAMES = [
    'cantidad', 'cant', 'qty', 'quantity', 'piezas', 'cantidadsolicitada',
    'cantidadpedida', 'solicitado', 'pedido', 'cantidadpieza',
]
UNIT_COLUMN_NAMES = [
    'unidad', 'unid', 'uom', 'um', 'medida', 'presentacion', 'unidadmedida',
    'unidaddemedida', 'ump', 'unidaddemed',
]

# Palabras que delatan una fila de encabezado
HEADER_KEYWORDS = [
    'descrip', 'codigo', 'clave', 'cantidad', 'cant', 'unidad', 'unid',
    'precio', 'marca', 'sku', 'reng', 'partida', 'articulo', 'producto',
    'modelo', 'ref', 'item',
]

# Tokens comunes de unidades (para la heurística sin encabezado)
UNIT_TOKENS = {
    'PZ', 'PZA', 'PIEZA', 'PIEZAS', 'M', 'MT', 'MTS', 'METRO', 'METROS',
    'KG', 'G', 'GR', 'L', 'LT', 'ML', 'CAJA', 'CJA', 'PAQ', 'PAQUETE',
    'JGO', 'JUEGO', 'ROLLO', 'ROLLOS', 'PAR', 'PARES', 'BULTO', 'SACO',
    'GAL', 'GALON', 'TRAMO', 'TRAMOS', 'UND', 'UN', 'EA', 'KIT',
}


# ---------------------------------------------------------------------------
# Normalización
# ---------------------------------------------------------------------------
def strip_accents(text):
    """
    Quita acentos y diacríticos latinos de cualquier texto (mayúsculas o
    minúsculas). Usa NFKD para descomponer caracteres como 'é' -> 'e' + tilde
    y luego elimina las marcas combinantes. La 'ñ' se trata aparte porque en
    español es una letra distinta de 'n', no una 'n' acentuada.
    """
    if text is None:
        return ""
    text = str(text)
    text = text.replace('Ñ', 'N').replace('ñ', 'n')
    nfkd = unicodedata.normalize('NFKD', text)
    return ''.join(c for c in nfkd if not unicodedata.combining(c))


def normalize(text):
    if not text:
        return ""
    text = strip_accents(text).upper().strip()
    text = re.sub(r'[^\w\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    tokens = [SYNONYM_MAP.get(t, t) for t in text.split()]
    return ' '.join(t for t in tokens if t)


def normalize_code(code):
    """Normaliza un código de producto: mayúsculas, sin espacios ni símbolos."""
    if code is None:
        return ""
    code = str(code).strip().upper()
    code = re.sub(r'[\s\-\.]+', '', code)
    return code


def extract_keywords(norm_text, min_len=3):
    """Extrae palabras significativas eliminando stopwords y palabras muy cortas."""
    return [w for w in norm_text.split() if w not in STOPWORDS and len(w) >= min_len]


# ---------------------------------------------------------------------------
# Detección y extracción del código del producto
# ---------------------------------------------------------------------------
def detect_code_column(rows):
    """Busca una columna que parezca contener el código del producto."""
    if not rows:
        return None
    columns = list(rows[0].keys())
    for col in columns:
        col_lower = col.lower().replace(' ', '').replace('_', '')
        for candidate in CODE_COLUMN_NAMES:
            cand_clean = candidate.replace('_', '')
            if col_lower == cand_clean or cand_clean in col_lower:
                return col
    return None


def extract_code_from_text(text):
    """
    Intenta extraer un código de producto del texto.
    Retorna el primer candidato razonable encontrado.
    """
    if not text:
        return None
    text = str(text).strip()

    m = re.search(r'\b(\d{5,})\b', text)
    if m:
        return m.group(1)

    candidates = re.findall(r'\b([A-Z]+[\-\.]?\d+[A-Z0-9\-\.]*)\b', text.upper())
    if candidates:
        excluded = {'3P', '2P', '1P', 'KW', 'HP', 'HZ', 'MM', 'PZ',
                    'V', 'A', '24V', '110V', '220V', '440V', '60HZ', '50HZ'}
        for c in candidates:
            clean = re.sub(r'[\-\.]', '', c)
            if clean not in excluded and len(clean) >= 4:
                return c

    return None


# ---------------------------------------------------------------------------
# Construcción de entrada de catálogo
# ---------------------------------------------------------------------------
def build_catalog_entry(row):
    name = row.get('DescCortaArt', '') or ''
    norm = normalize(name)
    code = row.get('CodigoArt', '') or ''
    return {
        'code':      code,
        'norm_code': normalize_code(code),
        'name':      name,
        'price':     str(row.get('Precio', '') or ''),
        'norm':      norm,
        'norm_set':  set(norm.split()),
    }


# ===========================================================================
# CATÁLOGO EN MEMORIA (caché) + ÍNDICE INVERTIDO
# ===========================================================================
_CACHE = {
    "catalog":   None,
    "by_code":   {},
    "by_norm":   {},
    "inverted":  {},
    "loaded_at": 0,
}
_CACHE_TTL = 6 * 3600

_supabase = None


def _get_supabase():
    """Crea el cliente de Supabase una sola vez y lo reutiliza."""
    global _supabase
    if _supabase is None:
        from supabase import create_client  # import perezoso
        _supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
    return _supabase


def _new_supabase(service=False):
    """Crea un cliente nuevo de Supabase (para hilos en segundo plano).
    Con service=True usa la service_role key para poder ESCRIBIR jobs/job_items."""
    from supabase import create_client
    key = SUPABASE_SERVICE_KEY if service else SUPABASE_KEY
    return create_client(SUPABASE_URL, key)


def _now_iso():
    return datetime.now(timezone.utc).isoformat()


def _fetch_all_products():
    """Trae TODO el catálogo paginando (PostgREST limita a 1000 filas)."""
    PAGE_SIZE = 1000
    rows = []
    offset = 0
    supabase = _get_supabase()
    while True:
        resp = supabase.table("products").select(
            "CodigoArt, DescCortaArt, Precio"
        ).range(offset, offset + PAGE_SIZE - 1).execute()

        page = resp.data or []
        if not page:
            break

        rows.extend(page)

        if len(page) < PAGE_SIZE:
            break

        offset += PAGE_SIZE

        if offset > 100000:
            break

    return rows


def _build_inverted(catalog):
    """Índice invertido: palabra clave -> conjunto de índices de productos."""
    inv = {}
    for i, entry in enumerate(catalog):
        for tok in entry['norm_set']:
            if len(tok) >= 3 and tok not in STOPWORDS:
                inv.setdefault(tok, set()).add(i)
    return inv


def load_catalog(force=False):
    """Carga (o recarga) el catálogo en memoria y reconstruye los índices."""
    fresh = _CACHE["catalog"] is not None and (
        _CACHE_TTL is None or (time.time() - _CACHE["loaded_at"] < _CACHE_TTL)
    )
    if fresh and not force:
        return

    raw = _fetch_all_products()
    catalog = [build_catalog_entry(row) for row in raw]

    by_code = {e['norm_code']: e for e in catalog if e['norm_code']}

    by_norm = {}
    for e in catalog:
        if e['norm'] and e['norm'] not in by_norm:
            by_norm[e['norm']] = e

    _CACHE.update({
        "catalog":   catalog,
        "by_code":   by_code,
        "by_norm":   by_norm,
        "inverted":  _build_inverted(catalog),
        "loaded_at": time.time(),
    })


def get_catalog():
    """Devuelve la caché, cargándola de forma perezosa si aún no existe."""
    if _CACHE["catalog"] is None or (
        _CACHE_TTL is not None and time.time() - _CACHE["loaded_at"] >= _CACHE_TTL
    ):
        load_catalog()
    return _CACHE


def candidates(query, top_n=100):
    """Pre-filtro: hasta `top_n` productos que comparten palabras clave."""
    cache = get_catalog()
    kws = extract_keywords(normalize(query))
    if not kws:
        return []

    inv = cache["inverted"]
    catalog = cache["catalog"]
    cnt = Counter()
    for kw in kws:
        for idx in inv.get(kw, ()):
            cnt[idx] += 1

    if not cnt:
        return []

    return [catalog[i] for i, _ in cnt.most_common(top_n)]


# ---------------------------------------------------------------------------
# PASADA 0 — Búsqueda directa por código (en memoria)
# ---------------------------------------------------------------------------
def code_lookup(code, by_code):
    if not code:
        return None
    norm = normalize_code(code)
    if not norm:
        return None
    return by_code.get(norm)


# ---------------------------------------------------------------------------
# PASADA 2 — Fuzzy Match (SequenceMatcher + Jaccard) sobre candidatos
# ---------------------------------------------------------------------------
def fuzzy_match(query, entries):
    if not entries:
        return None, 0.0

    norm_q  = normalize(query)
    q_words = set(norm_q.split())

    scored = []
    for entry in entries:
        norm_e  = entry['norm']
        e_words = entry['norm_set']

        seq     = difflib.SequenceMatcher(None, norm_q, norm_e).ratio()
        overlap = (len(q_words & e_words) / max(len(q_words), len(e_words))
                   if q_words and e_words else 0.0)
        combined = 0.5 * seq + 0.5 * overlap
        scored.append((combined, entry))

    scored.sort(key=lambda x: x[0], reverse=True)
    best_score, best_entry = scored[0]
    return best_entry, round(best_score, 3)


# ---------------------------------------------------------------------------
# PASADA 3 — Intersección de palabras clave sobre candidatos
# ---------------------------------------------------------------------------
def keyword_ilike_match(query, entries):
    norm_q   = normalize(query)
    keywords = extract_keywords(norm_q)

    if not keywords or not entries:
        return None, 0, 0

    code_counter = Counter()
    for kw in keywords:
        for entry in entries:
            if kw in entry['norm']:
                code_counter[entry['code']] += 1

    if not code_counter:
        return None, 0, 0

    cand_by_code = {e['code']: e for e in entries if e['code']}

    total_keywords = len(keywords)
    max_count = code_counter.most_common(1)[0][1]
    winners = [c for c, count in code_counter.items() if count == max_count]

    if len(winners) == 1:
        best_code = winners[0]
    else:
        best_code = max(
            winners,
            key=lambda c: difflib.SequenceMatcher(
                None, norm_q, cand_by_code[c]['norm']
            ).ratio()
        )

    best_entry = cand_by_code.get(best_code)
    if best_entry is None:
        return None, 0, 0

    return {
        'code':  best_entry['code'],
        'name':  best_entry['name'],
        'price': best_entry['price'],
        'norm':  best_entry['norm'],
    }, max_count, total_keywords


# ---------------------------------------------------------------------------
# Lógica principal de matching (tres pasadas)
# ---------------------------------------------------------------------------
def match_product(query, code):
    cache    = get_catalog()
    by_code  = cache["by_code"]
    by_norm  = cache["by_norm"]

    # --- Pasada 0: Código del producto ---
    code_to_try = code
    if not code_to_try and query:
        code_to_try = extract_code_from_text(query)

    if code_to_try:
        match0 = code_lookup(code_to_try, by_code)
        if match0:
            return {
                "codigo":            match0['code'],
                "nombre_catalogo":   match0['name'],
                "precio":            match0['price'],
                "confianza":         1.0,
                "requiere_revision": False,
                "metodo":            "codigo",
            }

    if not query or not query.strip():
        return {
            "codigo":            "NO ENCONTRADO",
            "nombre_catalogo":   "NO ENCONTRADO",
            "precio":            "",
            "confianza":         0.0,
            "requiere_revision": True,
            "metodo":            "ninguno",
        }

    # --- Pasada 1: Match exacto por descripción normalizada ---
    norm_query = normalize(query)
    if norm_query and norm_query in by_norm:
        match_exact = by_norm[norm_query]
        return {
            "codigo":            match_exact['code'],
            "nombre_catalogo":   match_exact['name'],
            "precio":            match_exact['price'],
            "confianza":         1.0,
            "requiere_revision": False,
            "metodo":            "exacto",
        }

    # --- Pre-filtro: candidatos del índice invertido ---
    cand = candidates(query)
    if not cand:
        return {
            "codigo":            "NO ENCONTRADO",
            "nombre_catalogo":   "NO ENCONTRADO",
            "precio":            "",
            "confianza":         0.0,
            "requiere_revision": True,
            "metodo":            "ninguno",
        }

    # --- Pasada 2: Fuzzy ---
    match1, score1 = fuzzy_match(query, cand)
    if match1 and score1 >= 0.70:
        return {
            "codigo":            match1['code'],
            "nombre_catalogo":   match1['name'],
            "precio":            match1['price'],
            "confianza":         score1,
            "requiere_revision": False,
            "metodo":            "fuzzy",
        }

    # --- Pasada 3: Intersección de keywords sobre candidatos ---
    match2, matched_kw, total_kw = keyword_ilike_match(query, cand)

    if match2 is None or matched_kw == 0:
        return {
            "codigo":            "NO ENCONTRADO",
            "nombre_catalogo":   "NO ENCONTRADO",
            "precio":            "",
            "confianza":         0.0,
            "requiere_revision": True,
            "metodo":            "ninguno",
        }

    if total_kw >= 5:
        min_required = 3
    elif total_kw >= 3:
        min_required = 2
    else:
        min_required = 1

    if matched_kw < min_required:
        return {
            "codigo":            "NO ENCONTRADO",
            "nombre_catalogo":   "NO ENCONTRADO",
            "precio":            "",
            "confianza":         0.0,
            "requiere_revision": True,
            "metodo":            "ninguno",
        }

    ratio      = matched_kw / total_kw
    confidence = round(0.40 + ratio * 0.29, 3)

    return {
        "codigo":            match2['code'],
        "nombre_catalogo":   match2['name'],
        "precio":            match2['price'],
        "confianza":         confidence,
        "requiere_revision": True,
        "metodo":            "keyword",
    }


# ---------------------------------------------------------------------------
# Detección automática de columna de descripción
# ---------------------------------------------------------------------------
def detect_description_column(rows):
    if not rows:
        return None
    priority_names = [
        'descripcion', 'description', 'producto', 'product',
        'item', 'articulo', 'article', 'nombre', 'name',
        'desc', 'material', 'concepto', 'detalle'
    ]
    columns = list(rows[0].keys())
    for col in columns:
        if any(p in col.lower() for p in priority_names):
            return col
    best_col, best_avg = None, 0
    for col in columns:
        values  = [str(r.get(col, '') or '') for r in rows[:10]]
        avg_len = sum(len(v) for v in values) / max(len(values), 1)
        if avg_len > best_avg:
            best_avg = avg_len
            best_col = col
    return best_col


def _detect_named_column(rows, names):
    """Detección genérica de columna por nombre de encabezado (cant/unidad)."""
    if not rows:
        return None
    cols = list(rows[0].keys())
    norm_cols = [(col, strip_accents(str(col).lower()).replace(' ', '').replace('_', ''))
                 for col in cols]
    # Coincidencia exacta primero
    for col, c in norm_cols:
        for cand in names:
            cc = strip_accents(cand.lower()).replace(' ', '').replace('_', '')
            if c == cc:
                return col
    # Coincidencia por substring solo para nombres de >=4 chars (evita 'um','uom')
    for col, c in norm_cols:
        for cand in names:
            cc = strip_accents(cand.lower()).replace(' ', '').replace('_', '')
            if len(cc) >= 4 and cc in c:
                return col
    return None


# ===========================================================================
# FASE 2 — Validación semántica con LLM (deltas-only) — SIN CAMBIOS
# ===========================================================================
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
ANTHROPIC_MODEL   = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-6")
LLM_BATCH_SIZE    = int(os.environ.get("LLM_BATCH_SIZE", "40"))
LLM_MAX_WORKERS   = int(os.environ.get("LLM_MAX_WORKERS", "4"))

_anthropic_client = None


def _get_anthropic():
    global _anthropic_client
    if _anthropic_client is None:
        import anthropic  # import perezoso
        _anthropic_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    return _anthropic_client


def _join_text(msg):
    """Une los bloques de texto de una respuesta del SDK de Anthropic."""
    return "".join(
        getattr(b, "text", "") for b in msg.content
        if getattr(b, "type", "") == "text"
    )


EVALUATOR_SYSTEM = (
    "Eres un evaluador semantico de productos industriales en Mexico. Recibiras "
    "un arreglo JSON de lineas, cada una con: idx, descripcion_original, "
    "nombre_catalogo y metodo. Valida si nombre_catalogo corresponde "
    "semanticamente a descripcion_original y devuelve SOLO las correcciones.\n\n"
    "CRITERIOS DE CONFIANZA:\n"
    "- Mismo producto escrito distinto (espacios, mayusculas, orden) -> 1.0\n"
    "- Mismo producto con abreviaturas conocidas o typos menores -> 0.90-0.99\n"
    "- Probablemente el mismo pero con duda razonable -> 0.70-0.89\n"
    "- Relacionado pero con diferencias importantes (modelo, capacidad, voltaje) -> 0.50-0.69\n"
    "- No relacionados / completamente diferentes -> 0.0-0.49\n\n"
    "ABREVIACIONES (no penalices por estas): P/=Para, C/=Con, S/=Sin, Jgo=Juego, "
    "Mca=Marca, No.=Numero, Pza/Pz=Pieza, Term=Termomagnetico, Interr=Interruptor, "
    "Volt/V=Voltaje, Amp/A=Amperes, Polos/P=Polos.\n\n"
    "REGLA NO ENCONTRADO: si nombre_catalogo ya dice \"NO ENCONTRADO\", o si el "
    "producto no tiene NINGUNA relacion con la descripcion, pon en esa linea "
    "codigo=\"NO ENCONTRADO\", nombre_catalogo=\"NO ENCONTRADO\", confianza=0.0.\n\n"
    "REGLA DE REVISION: requiere_revision = true si y solo si confianza < 0.90.\n\n"
    "SALIDA (CRITICO): responde UNICAMENTE con un arreglo JSON valido, sin texto "
    "ni markdown, empezando con '[' y terminando con ']'. Para cada linea de "
    "entrada incluye exactamente: {\"idx\": <n>, \"confianza\": <0.0-1.0>, "
    "\"requiere_revision\": <bool>}. Incluye ademas \"codigo\" y \"nombre_catalogo\" "
    "SOLO si la cambias a \"NO ENCONTRADO\". No agregues otros campos ni lineas "
    "que no esten en la entrada."
)


def _parse_deltas(text):
    """Extrae el arreglo JSON de la respuesta del modelo de forma robusta."""
    s = (text or "").strip()
    a, b = s.find("["), s.rfind("]")
    if a == -1 or b == -1 or b < a:
        return []
    try:
        return json.loads(s[a:b + 1])
    except Exception:
        return []


def _eval_batch(batch):
    """Evalúa un lote de líneas ambiguas y devuelve la lista de deltas."""
    payload = [
        {
            "idx":                  r["_idx"],
            "descripcion_original": r["descripcion_original"],
            "nombre_catalogo":      r["nombre_catalogo"],
            "metodo":               r["metodo"],
        }
        for r in batch
    ]
    try:
        client = _get_anthropic()
        msg = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=4096,
            system=EVALUATOR_SYSTEM,
            messages=[{"role": "user",
                       "content": json.dumps(payload, ensure_ascii=False)}],
        )
        return _parse_deltas(_join_text(msg))
    except Exception:
        return []


def evaluate_with_llm(resultados):
    """Validación semántica del LLM SOLO a las líneas fuzzy/keyword."""
    if not ANTHROPIC_API_KEY:
        return resultados

    ambiguous = []
    for i, r in enumerate(resultados):
        if r.get("metodo") in ("fuzzy", "keyword"):
            r["_idx"] = i
            ambiguous.append(r)

    if not ambiguous:
        return resultados

    batches = [ambiguous[k:k + LLM_BATCH_SIZE]
               for k in range(0, len(ambiguous), LLM_BATCH_SIZE)]

    deltas = []
    workers = max(1, min(LLM_MAX_WORKERS, len(batches)))
    with ThreadPoolExecutor(max_workers=workers) as ex:
        for res in ex.map(_eval_batch, batches):
            deltas.extend(res or [])

    for d in deltas:
        if not isinstance(d, dict):
            continue
        i = d.get("idx")
        if not isinstance(i, int) or not (0 <= i < len(resultados)):
            continue
        r = resultados[i]
        if d.get("codigo") == "NO ENCONTRADO" or d.get("nombre_catalogo") == "NO ENCONTRADO":
            r["codigo"]            = "NO ENCONTRADO"
            r["nombre_catalogo"]   = "NO ENCONTRADO"
            r["precio"]            = ""
            r["confianza"]         = 0.0
            r["requiere_revision"] = True
            continue
        if "confianza" in d:
            try:
                conf = round(float(d["confianza"]), 3)
            except (TypeError, ValueError):
                continue
            r["confianza"]         = conf
            r["requiere_revision"] = bool(conf < 0.90)

    for r in resultados:
        r.pop("_idx", None)

    return resultados


# ===========================================================================
# FASE 0 (NUEVA) — EXTRACCIÓN: archivo -> JSON canónico
# ---------------------------------------------------------------------------
# Reemplaza a la Edge Function de Supabase y a la decodificación del navegador.
# Salida idéntica a la que Bolt ya espera: { "data": [ {IEST-01, Codigo,
# Descripcion, Unid, Cant} ] }.
# ===========================================================================
DOCLING_OCR_URL = os.environ.get(
    "DOCLING_OCR_URL",
    "https://cotizaciones-docling-production.up.railway.app/ocr",
)
DOCLING_TIMEOUT = int(os.environ.get("DOCLING_TIMEOUT", "180"))

# Prompt de extracción (portado de la Edge Function, comportamiento probado).
EXTRACTOR_SYSTEM = """You are a procurement specialist extracting products from invoices, quotes, purchase orders, and product lists in Spanish.

CRITICAL RULES:
1. Extract ONLY product line items — rows that represent actual products, materials, or items to be purchased/sold
2. HEADER DETECTION: The document may or may not have a header row. If the first row looks like column names (e.g., "Descripcion", "Cantidad", "Producto", "Articulo", "Unidad"), skip it. If the first row looks like an actual product (e.g., a product code, part number, or product description), include it — DO NOT skip it.
3. Each row with product data must be included, including the first row if it is a product
4. Ignore: footers, subtotals, totals, taxes, IVA, transport charges, payment terms, dates, signatures, company info
5. Each product MUST have a description
6. Extract quantities exactly as shown. If no quantity is found or the value is 0, default to 1
7. If no unit is found, use "PZ" (piezas)
8. Never create or invent products not present in the document
9. Preserve original descriptions exactly as written
10. For handwritten lists: read carefully and extract every item listed
11. For tab-separated or CSV-like text: treat each row as a separate product line
12. PRODUCT CODE EXTRACTION: If the document has a column with product codes, SKUs, part numbers, or item references (column headers like 'Codigo', 'Clave', 'SKU', 'CodigoArt', 'No. Parte', 'Cve', 'Referencia', 'Item', 'Ref'), extract that code into the 'Codigo' field of each row. If no code is present for a row, set 'Codigo' to an empty string. NEVER invent or guess product codes — only extract what is explicitly written in the document.

RETURN FORMAT — ONLY a valid JSON array, NO markdown, NO backticks, NO explanation:

[
  {
    "Codigo": "exact product code as written, or empty string if absent",
    "Descripcion": "exact product description",
    "Cantidad": number_value,
    "Unidad": "unit abbreviation"
  }
]

WHAT TO IGNORE:
- "Subtotal", "IVA", "Total", "TOTAL", "Gran Total"
- "Documento generado", "Fecha", "Referencia", "Folio"
- "Condiciones de pago", "Transporte", "Agente", "Vendedor"
- Rows with only numeric totals or calculations
- Company name, address, RFC, contact info

IMPORTANT: When in doubt about whether a row is a product or a header, include it as a product. It is better to include an extra row than to miss a real product."""

IMG_MEDIA_TYPES = {
    'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
    'webp': 'image/webp', 'gif': 'image/gif',
}


# ------------------------- Helpers de salida canónica -----------------------
def _format_canonical(rows):
    """Convierte filas intermedias {Codigo, Descripcion, Cantidad, Unidad} en el
    JSON canónico final con IEST-01 secuencial. Igual que la Edge Function."""
    out = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        desc = str(row.get('Descripcion') or '').strip()
        if not desc:
            continue

        cant = 1.0
        rc = row.get('Cantidad')
        if rc is not None and str(rc).strip() != '':
            try:
                x = float(str(rc).replace(',', '.'))
                if x > 0 and math.isfinite(x):
                    cant = x
            except (TypeError, ValueError):
                pass

        unid = (str(row.get('Unidad') or 'PZ').strip()[:10]) or 'PZ'
        cant_str = str(int(cant)) if float(cant).is_integer() else str(cant)

        out.append({
            'Codigo':      str(row.get('Codigo') or '').strip(),
            'Descripcion': desc,
            'Unid':        unid,
            'Cant':        cant_str,
        })

    return [
        {
            'IEST-01':     str(i + 1),
            'Codigo':      r['Codigo'],
            'Descripcion': r['Descripcion'],
            'Unid':        r['Unid'],
            'Cant':        r['Cant'],
        }
        for i, r in enumerate(out)
    ]


def _parse_extracted_json(text):
    """Parsea de forma robusta el arreglo JSON devuelto por el LLM de extracción."""
    s = (text or "").strip().replace('```json', '').replace('```', '').replace('`', '').strip()
    a, b = s.find('['), s.rfind(']')
    if a == -1 or b == -1 or b < a:
        return []
    frag = s[a:b + 1]
    for attempt in (frag, frag.replace('\n', ' ').replace('\r', ' ').replace(',]', ']').replace(',}', '}')):
        try:
            v = json.loads(attempt)
            if isinstance(v, list):
                return v
        except Exception:
            pass
    # Reparación de arreglo truncado: corta en el último '}' y cierra.
    last = frag.rfind('}')
    if last != -1:
        try:
            v = json.loads(frag[:last + 1] + ']')
            if isinstance(v, list):
                return v
        except Exception:
            pass
    return []


# ------------------------- Camino con LLM (texto) ---------------------------
def _looks_like_header(text_or_cells):
    """¿La primera fila/línea parece encabezado? (>=2 keywords y sin números largos)"""
    if isinstance(text_or_cells, (list, tuple)):
        line = ' '.join(str(c) for c in text_or_cells)
    else:
        line = str(text_or_cells)
    low = strip_accents(line.lower())
    count = sum(1 for k in HEADER_KEYWORDS if strip_accents(k) in low)
    return count >= 2 and not re.search(r'\d{4,}', line)


def _extract_chunk_llm(chunk_text, source_type):
    try:
        client = _get_anthropic()
        msg = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=8192,
            system=EXTRACTOR_SYSTEM,
            messages=[{
                "role": "user",
                "content": (
                    f"Extract all product line items from this {source_type} document. "
                    f"Be thorough and extract every product row:\n\n{chunk_text[:80000]}"
                ),
            }],
        )
        return _parse_extracted_json(_join_text(msg))
    except Exception:
        return []


def extract_text_llm(text, source_type):
    """Trocea el texto y llama al LLM EN PARALELO (no en serie como antes)."""
    lines = [l.rstrip() for l in str(text).split('\n') if l.strip() != '']
    if not lines:
        return []

    has_header = _looks_like_header(lines[0])
    header = lines[0] if has_header else None
    data_lines = lines[1:] if has_header else lines

    CHUNK = LLM_BATCH_SIZE
    chunks = [data_lines[i:i + CHUNK] for i in range(0, len(data_lines), CHUNK)]
    payloads = [((header + '\n') if header else '') + '\n'.join(c) for c in chunks if c]
    if not payloads:
        payloads = [str(text)]

    results = []
    workers = max(1, min(LLM_MAX_WORKERS, len(payloads)))
    with ThreadPoolExecutor(max_workers=workers) as ex:
        for rows in ex.map(lambda ct: _extract_chunk_llm(ct, source_type), payloads):
            results.extend(rows or [])
    return results


# ------------------------- Camino con LLM (imagen / PDF doc) ----------------
def extract_image_llm(content, media_type):
    """Foto / escaneo / manuscrito -> Claude visión directo."""
    try:
        b64 = base64.b64encode(content).decode('ascii')
        client = _get_anthropic()
        msg = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=8192,
            system=EXTRACTOR_SYSTEM,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image", "source": {
                        "type": "base64", "media_type": media_type, "data": b64}},
                    {"type": "text", "text": (
                        "Extract all product line items from this image. This may be a "
                        "photo of a spreadsheet, a printed form, or a handwritten list. "
                        "Be thorough and extract every product row you can read.")},
                ],
            }],
        )
        return _parse_extracted_json(_join_text(msg))
    except Exception:
        return []


def extract_pdf_document_llm(content):
    """Fallback: manda el PDF como documento a Claude (si Docling no devolvió nada)."""
    try:
        b64 = base64.b64encode(content).decode('ascii')
        client = _get_anthropic()
        msg = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=8192,
            system=EXTRACTOR_SYSTEM,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "document", "source": {
                        "type": "base64", "media_type": "application/pdf", "data": b64}},
                    {"type": "text", "text": (
                        "Extract all product line items from this PDF. It may be a "
                        "scanned document. Be thorough and extract every product row.")},
                ],
            }],
        )
        return _parse_extracted_json(_join_text(msg))
    except Exception:
        return []


# ------------------------- Docling (servicio HTTP) --------------------------
def _pick_first(d, keys):
    for k in keys:
        if isinstance(d, dict) and d.get(k) not in (None, ''):
            return d[k]
    return None


def _find_rows_array(raw):
    if isinstance(raw, list):
        return raw
    if not isinstance(raw, dict):
        return None
    for k in ['rows', 'lines', 'items', 'data', 'products', 'productos',
              'partidas', 'records', 'extracted', 'result', 'results', 'output']:
        if isinstance(raw.get(k), list):
            return raw[k]
    for v in raw.values():
        if isinstance(v, list) and v and isinstance(v[0], dict):
            return v
    return None


def _map_docling(raw):
    arr = _find_rows_array(raw)
    if not arr:
        return []
    out = []
    for row in arr:
        if not isinstance(row, dict):
            continue
        desc = str(_pick_first(row, [
            'Descripcion', 'descripcion', 'description', 'desc', 'nombre', 'name',
            'producto', 'item', 'product', 'product_name', 'itemName',
            'descripcion_producto']) or '').strip()
        if not desc:
            continue
        codigo = str(_pick_first(row, [
            'Codigo', 'codigo', 'code', 'sku', 'clave', 'item_code', 'product_code',
            'itemCode', 'productCode', 'codigo_articulo', 'articulo', 'no_parte',
            'noParte', 'ref', 'referencia']) or '').strip()
        unidad = _pick_first(row, [
            'Unid', 'unid', 'unidad', 'unit', 'uom', 'unit_of_measure',
            'unidad_medida', 'ump'])
        cant = _pick_first(row, [
            'Cant', 'cant', 'cantidad', 'qty', 'quantity', 'count', 'pieces'])
        out.append({
            'Codigo': codigo,
            'Descripcion': desc,
            'Cantidad': cant if cant is not None else '',
            'Unidad': unidad if unidad is not None else '',
        })
    return out


def _docling_extract(content, filename):
    """Manda el archivo al servicio Docling de Railway y mapea su respuesta."""
    try:
        import requests
        resp = requests.post(
            DOCLING_OCR_URL,
            files={'file': (filename, content)},
            timeout=DOCLING_TIMEOUT,
        )
        if resp.status_code != 200:
            return []
        raw = None
        try:
            raw = resp.json()
        except Exception:
            txt = resp.text.strip()
            if txt[:1] in ('[', '{'):
                raw = json.loads(txt)
        return _map_docling(raw) if raw is not None else []
    except Exception:
        return []


# ------------------------- PDF -----------------------------------------------
def extract_pdf(content, filename):
    """PDF con texto -> extracción local + LLM. PDF escaneado -> Docling -> Claude."""
    text = ''
    npages = 0
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            npages = len(pdf.pages)
            parts = []
            for p in pdf.pages:
                t = p.extract_text() or ''
                if t.strip():
                    parts.append(t)
            text = '\n\n'.join(parts).strip()
    except Exception:
        text = ''

    avg = (len(text) / npages) if npages else len(text)
    if len(text) >= 20 and avg >= 15:
        return extract_text_llm(text, 'pdf')

    # Escaneado / sin capa de texto -> Docling
    rows = _docling_extract(content, filename or 'file.pdf')
    if rows:
        return rows
    # Último recurso: el PDF como documento a Claude
    return extract_pdf_document_llm(content)


# ------------------------- DOCX ----------------------------------------------
def extract_docx_llm(content):
    text = ''
    try:
        import docx
        d = docx.Document(io.BytesIO(content))
        parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append('\t'.join(cells))
        text = '\n'.join(parts).strip()
    except Exception:
        text = ''
    if not text:
        return _docling_extract(content, 'file.docx')
    return extract_text_llm(text, 'docx')


# ------------------------- Estructurados (xlsx/xls/csv) ----------------------
def _df_to_grid(df):
    grid = []
    for row in df.values.tolist():
        cells = []
        for c in row:
            if c is None:
                cells.append('')
            else:
                s = str(c)
                if s.strip().lower() == 'nan':
                    s = ''
                cells.append(s.strip())
        if any(cell != '' for cell in cells):
            grid.append(cells)
    return grid


def _read_grids(content, ext):
    """Lee xlsx/xls/csv como rejillas de strings (todas las filas, todas las hojas)."""
    import pandas as pd
    grids = []
    bio = io.BytesIO(content)

    if ext == 'csv':
        df = None
        for enc in ('utf-8-sig', 'utf-8', 'latin-1'):
            try:
                bio.seek(0)
                df = pd.read_csv(bio, header=None, dtype=str, keep_default_na=False,
                                 sep=None, engine='python', encoding=enc)
                break
            except Exception:
                df = None
        if df is None:
            bio.seek(0)
            df = pd.read_csv(bio, header=None, dtype=str, keep_default_na=False,
                             encoding='latin-1')
        grids.append(_df_to_grid(df))
    else:
        engine = 'xlrd' if ext == 'xls' else 'openpyxl'
        bio.seek(0)
        sheets = pd.read_excel(bio, header=None, dtype=str, sheet_name=None, engine=engine)
        for _name, df in sheets.items():
            grids.append(_df_to_grid(df))
    return grids


def _rows_from_header_grid(grid):
    """Rejilla CON encabezado -> filas intermedias, detectando columnas por nombre."""
    header = grid[0]
    keys, seen = [], {}
    for j, h in enumerate(header):
        k = (h or '').strip() or f'col_{j}'
        if k in seen:
            seen[k] += 1
            k = f'{k}_{seen[k]}'
        else:
            seen[k] = 0
        keys.append(k)

    dict_rows = []
    for r in grid[1:]:
        d = {keys[j]: (r[j] if j < len(r) else '') for j in range(len(keys))}
        dict_rows.append(d)

    if not dict_rows:
        return []

    desc_col = detect_description_column(dict_rows)
    code_col = detect_code_column(dict_rows)
    qty_col  = _detect_named_column(dict_rows, QTY_COLUMN_NAMES)
    unit_col = _detect_named_column(dict_rows, UNIT_COLUMN_NAMES)

    out = []
    for d in dict_rows:
        desc = str(d.get(desc_col, '') if desc_col else '').strip()
        if not desc:
            continue
        out.append({
            'Codigo':      str(d.get(code_col, '') if code_col else '').strip(),
            'Descripcion': desc,
            'Cantidad':    d.get(qty_col, '') if qty_col else '',
            'Unidad':      d.get(unit_col, '') if unit_col else '',
        })
    return out


def _is_code_like(s):
    s = str(s).strip().upper()
    if re.fullmatch(r'\d{5,}', s):
        return True
    if re.fullmatch(r'[A-Z]+[\-\.]?\d+[A-Z0-9\-\.]*', s) and len(re.sub(r'[\-\.]', '', s)) >= 4:
        return True
    return False


def _is_unit_like(s):
    s = strip_accents(str(s).strip().upper())
    if s in UNIT_TOKENS:
        return True
    return bool(s) and len(s) <= 4 and s.isalpha()


def _rows_from_headerless_grid(grid):
    """Rejilla SIN encabezado -> heurística por contenido de columnas."""
    ncols = max((len(r) for r in grid), default=0)
    if ncols == 0:
        return []

    def cell(r, j):
        return r[j] if j < len(r) else ''

    def numeric_ratio(j):
        vals = [cell(r, j) for r in grid if cell(r, j) != '']
        if not vals:
            return 0.0
        n = 0
        for v in vals:
            try:
                float(str(v).replace(',', '.'))
                n += 1
            except Exception:
                pass
        return n / len(vals)

    def avg_len(j):
        vals = [cell(r, j) for r in grid]
        return sum(len(v) for v in vals) / max(len(vals), 1)

    # Descripción: columna no numérica con texto más largo en promedio
    desc_idx, best = None, -1.0
    for j in range(ncols):
        if numeric_ratio(j) > 0.6:
            continue
        a = avg_len(j)
        if a > best:
            best, desc_idx = a, j
    if desc_idx is None:
        for j in range(ncols):
            a = avg_len(j)
            if a > best:
                best, desc_idx = a, j

    def ratio_for(j, pred):
        vals = [cell(r, j) for r in grid if cell(r, j) != '']
        if not vals:
            return 0.0
        return sum(1 for v in vals if pred(v)) / len(vals)

    code_idx, bestc = None, 0.4
    for j in range(ncols):
        if j == desc_idx:
            continue
        cr = ratio_for(j, _is_code_like)
        if cr > bestc:
            bestc, code_idx = cr, j

    def is_qty(v):
        try:
            x = float(str(v).replace(',', '.'))
            return 0 < x < 100000
        except Exception:
            return False

    qty_idx, bestq = None, 0.5
    for j in range(ncols):
        if j in (desc_idx, code_idx):
            continue
        q = ratio_for(j, is_qty)
        if q > bestq:
            bestq, qty_idx = q, j

    unit_idx, bestu = None, 0.5
    for j in range(ncols):
        if j in (desc_idx, code_idx, qty_idx):
            continue
        u = ratio_for(j, _is_unit_like)
        if u > bestu:
            bestu, unit_idx = u, j

    out = []
    for r in grid:
        desc = cell(r, desc_idx).strip() if desc_idx is not None else ''
        if not desc:
            continue
        out.append({
            'Codigo':      cell(r, code_idx).strip() if code_idx is not None else '',
            'Descripcion': desc,
            'Cantidad':    cell(r, qty_idx) if qty_idx is not None else '',
            'Unidad':      cell(r, unit_idx) if unit_idx is not None else '',
        })
    return out


def extract_structured(content, ext):
    """xlsx/xls/csv -> filas intermedias SIN LLM."""
    grids = _read_grids(content, ext)
    out = []
    for grid in grids:
        if not grid:
            continue
        if _looks_like_header(grid[0]):
            out.extend(_rows_from_header_grid(grid))
        else:
            out.extend(_rows_from_headerless_grid(grid))
    return out


# ------------------------- Router de extracción -----------------------------
def extract_file(content, filename):
    """Despacha por extensión y devuelve filas intermedias (sin formatear)."""
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in (filename or '') else ''

    if ext in ('xlsx', 'xlsm', 'xls', 'csv'):
        return extract_structured(content, ext)
    if ext in IMG_MEDIA_TYPES:
        return extract_image_llm(content, IMG_MEDIA_TYPES[ext])
    if ext == 'pdf':
        return extract_pdf(content, filename)
    if ext == 'docx':
        return extract_docx_llm(content)
    if ext == 'doc':
        return _docling_extract(content, filename or 'file.doc')
    if ext in ('txt', 'tsv'):
        return extract_text_llm(content.decode('utf-8', 'ignore'), ext)

    raise ValueError(f"Tipo de archivo no soportado: .{ext}")


# ===========================================================================
# FASE 3 (NUEVA) — MATCHING ASÍNCRONO POR LOTES (jobs + job_items)
# ---------------------------------------------------------------------------
# Bolt manda TODAS las filas a /match/start. Railway crea/actualiza el job,
# contesta de inmediato y procesa en segundo plano: por cada lote hace matching
# + validación LLM, inserta los renglones en job_items y actualiza jobs.progreso.
# Bolt queda libre; el usuario puede salir y volver a ver el avance leyendo
# las tablas jobs/job_items desde Supabase (que ya hace su frontend).
# ===========================================================================
def _line_from_item(item, desc_column, code_column):
    """Construye el renglón de resultado de UNA fila (matching, sin LLM)."""
    descripcion = (str(item.get(desc_column, '') or '').strip()
                   if desc_column else '')
    codigo_in = (str(item.get(code_column, '') or '').strip()
                 if code_column else '')
    if not descripcion and not codigo_in:
        return None

    cant = (item.get('Cant') or item.get('cant') or item.get('cantidad')
            or item.get('qty') or item.get('quantity') or '1')
    unid = (item.get('Unid') or item.get('unid') or item.get('unidad')
            or item.get('uom') or item.get('unit') or 'PZA')

    result = match_product(descripcion, codigo_in)
    return {
        "original_text":        descripcion,
        "descripcion_original": descripcion,
        "codigo_original":      codigo_in,
        "codigo":               result["codigo"],
        "nombre_catalogo":      result["nombre_catalogo"],
        "precio":               result["precio"],
        "confianza":            result["confianza"],
        "requiere_revision":    result["requiere_revision"],
        "metodo":               result["metodo"],
        "cant":                 cant,
        "unid":                 unid,
    }


def _num(x):
    """Convierte a número; devuelve None si no se puede."""
    if x is None:
        return None
    try:
        v = float(str(x).replace(',', '.'))
        return v if math.isfinite(v) else None
    except (TypeError, ValueError):
        return None


def _job_line_row(job_id, r):
    """Mapea un resultado de matching a una fila de la tabla job_lines."""
    no_match = (r.get('metodo') == 'ninguno'
                or str(r.get('codigo')) == 'NO ENCONTRADO')
    precio = None if no_match else _num(r.get('precio'))
    cantidad = _num(r.get('cant'))
    if cantidad is None or cantidad <= 0:
        cantidad = 1
    total = round(cantidad * precio, 2) if precio is not None else None
    unidad = str(r.get('unid') or '') or None
    return {
        'job_id':               job_id,
        'line_index':           r.get('_li', 0),
        'codigo_original':      (r.get('codigo_original') or None),
        'descripcion_original': (r.get('descripcion_original') or None),
        'unidad_original':      unidad,
        'cantidad':             cantidad,
        'producto_codigo':      None if no_match else (r.get('codigo') or None),
        'producto_descripcion': None if no_match else (r.get('nombre_catalogo') or None),
        'unidad_medida':        unidad,
        'precio_unitario':      precio,
        'confianza':            _num(r.get('confianza')),
        'requiere_revision':    bool(r.get('requiere_revision')),
        'total_linea':          total,
        'origen':               'sin_match' if no_match else 'auto',
        # 'estado' se queda con su default ('pendiente')
    }


def _run_match_job(referencia, rows):
    """Worker en segundo plano: procesa todas las filas por lotes, escribiendo
    el avance y los job_items en Supabase. Se ejecuta en un hilo daemon."""
    sb = _new_supabase(service=True)
    try:
        jr = sb.table('jobs').select('id').eq('referencia', referencia).limit(1).execute()
        if not jr.data:
            return
        job_id = jr.data[0]['id']

        get_catalog()  # asegurar catálogo en memoria
        desc_column = detect_description_column(rows)
        code_column = detect_code_column(rows)
        total = len(rows)

        sb.table('jobs').update({
            'status': 'matching', 'total_lineas': total,
            'progreso': 0, 'error': None, 'updated_at': _now_iso(),
        }).eq('id', job_id).execute()

        line_index = 0
        for start in range(0, total, JOB_CHUNK):
            chunk = rows[start:start + JOB_CHUNK]
            results = []
            for item in chunk:
                line = _line_from_item(item, desc_column, code_column)
                if line is not None:
                    line['_li'] = line_index
                    results.append(line)
                line_index += 1

            # Validación semántica LLM SOLO de las líneas fuzzy/keyword del lote.
            results = evaluate_with_llm(results)

            if results:
                payload = [_job_line_row(job_id, r) for r in results]
                try:
                    sb.table('job_lines').upsert(
                        payload, on_conflict='job_id,line_index'
                    ).execute()
                except Exception as e:
                    print(f"[job {referencia}] error guardando job_lines: {e}")

            processed = min(total, start + len(chunk))
            sb.table('jobs').update({
                'progreso': processed, 'updated_at': _now_iso(),
            }).eq('id', job_id).execute()

        sb.table('jobs').update({
            'status': 'completado', 'progreso': total, 'updated_at': _now_iso(),
        }).eq('id', job_id).execute()

    except Exception as e:
        try:
            sb.table('jobs').update({
                'status': 'error', 'error': str(e), 'updated_at': _now_iso(),
            }).eq('referencia', referencia).execute()
        except Exception:
            pass


# ---------------------------------------------------------------------------
# Helper de respuesta JSON
# ---------------------------------------------------------------------------
def _json(data, status=200):
    return Response(json.dumps(data, ensure_ascii=False), status=status,
                    mimetype='application/json')


# ===========================================================================
# Rutas Flask
# ===========================================================================
@app.route('/extract', methods=['POST'])
def extract_endpoint():
    """Archivo (multipart 'file') -> { "data": [ {IEST-01, Codigo, Descripcion,
    Unid, Cant} ] }. También acepta JSON {text, sourceType} o {imageBase64,
    mediaType} por compatibilidad con la Edge Function anterior."""
    try:
        f = request.files.get('file')

        if f is None:
            data = request.get_json(silent=True) or {}
            if data.get('text'):
                inter = extract_text_llm(str(data['text']), str(data.get('sourceType') or 'txt'))
                return _json({"data": _format_canonical(inter)})
            if data.get('imageBase64') and data.get('mediaType'):
                content = base64.b64decode(data['imageBase64'])
                inter = extract_image_llm(content, str(data['mediaType']))
                return _json({"data": _format_canonical(inter)})
            return _json({"error": "No se recibió archivo (campo 'file')"}, 400)

        filename = f.filename or 'upload'
        content = f.read()
        if not content:
            return _json({"error": "El archivo está vacío"}, 400)

        try:
            inter = extract_file(content, filename)
        except ValueError as ve:
            return _json({"error": str(ve)}, 400)

        rows = _format_canonical(inter)
        if not rows:
            # 200 con data vacía + aviso (igual que la Edge Function)
            return _json({"data": [], "error": "No se encontraron productos en el documento"})

        return _json({"data": rows})

    except Exception as e:
        return _json({"error": "Error interno", "message": str(e)}, 500)


@app.route('/match', methods=['POST'])
def match_products():
    try:
        data = request.get_json()
        if not data:
            return _json({"error": "No se recibieron datos"}, 400)

        client_products = (
            data.get('rows') or data.get('productos') or
            data.get('items') or data.get('articulos') or
            data.get('products') or []
        )

        if not client_products:
            return _json({"error": "No se encontraron productos"}, 400)

        desc_column = detect_description_column(client_products)
        code_column = detect_code_column(client_products)

        if not desc_column and not code_column:
            return _json(
                {"error": "No se pudo detectar columna de descripcion ni de codigo"}, 400)

        cache = get_catalog()
        if not cache["catalog"]:
            return _json({"error": "El catalogo esta vacio"}, 400)

        resultados = []
        for item in client_products:
            line = _line_from_item(item, desc_column, code_column)
            if line is not None:
                resultados.append(line)

        resultados = evaluate_with_llm(resultados)

        response_data = {
            "lines":               resultados,
            "total":               len(resultados),
            "requieren_revision":  sum(1 for r in resultados if r['requiere_revision']),
            "columna_descripcion": desc_column,
            "columna_codigo":      code_column,
        }
        return _json(response_data, 200)

    except Exception as e:
        return _json({"error": str(e)}, 500)


@app.route('/match/start', methods=['POST'])
def match_start():
    """Arranca el matching ASÍNCRONO. Crea/actualiza el job, lanza el worker en
    segundo plano y contesta de inmediato (202) con la referencia. Body:
    { referencia?, customerName?, rows: [...] }."""
    try:
        data = request.get_json()
        if not data:
            return _json({"error": "No se recibieron datos"}, 400)

        rows = (data.get('rows') or data.get('productos') or
                data.get('items') or data.get('articulos') or
                data.get('products') or [])
        if not rows:
            return _json({"error": "No se encontraron productos"}, 400)

        referencia = data.get('referencia') or data.get('reference')
        customer = data.get('customerName') or data.get('cliente') or ''

        sb = _new_supabase(service=True)

        if not referencia:
            referencia = f"QAI-{int(time.time() * 1000)}"

        # Asegura que el job EXISTA antes de arrancar el worker. Si Bolt ya lo
        # creo, lo actualiza; si no (o si hubo cualquier desajuste de referencia),
        # lo crea aqui mismo. Asi el procesamiento SIEMPRE arranca.
        existing = sb.table('jobs').select('id').eq(
            'referencia', referencia).limit(1).execute()
        if existing.data:
            sb.table('jobs').update({
                'cliente': customer or None, 'status': 'matching',
                'total_lineas': len(rows), 'progreso': 0, 'error': None,
                'updated_at': _now_iso(),
            }).eq('referencia', referencia).execute()
        else:
            sb.table('jobs').insert({
                'referencia': referencia, 'cliente': customer,
                'status': 'matching', 'total_lineas': len(rows), 'progreso': 0,
            }).execute()

        threading.Thread(
            target=_run_match_job, args=(referencia, rows), daemon=True
        ).start()

        return _json({
            "referencia": referencia, "status": "matching", "total": len(rows),
        }, 202)

    except Exception as e:
        return _json({"error": str(e)}, 500)


@app.route('/jobs/<referencia>', methods=['GET'])
def job_status(referencia):
    """Estado de un job por referencia (avance, total, status, error).
    Bolt también puede leer esto directo de Supabase; este endpoint es opcional."""
    try:
        sb = _new_supabase(service=True)
        jr = sb.table('jobs').select('*').eq(
            'referencia', referencia).limit(1).execute()
        if not jr.data:
            return _json({"error": "job no encontrado"}, 404)
        job = jr.data[0]
        try:
            ci = sb.table('job_lines').select(
                'id', count='exact').eq('job_id', job['id']).execute()
            job['items_count'] = getattr(ci, 'count', None)
        except Exception:
            job['items_count'] = None
        return _json(job)
    except Exception as e:
        return _json({"error": str(e)}, 500)


@app.route('/catalog/refresh', methods=['POST'])
def catalog_refresh():
    """Invalida y recarga la caché del catálogo."""
    try:
        load_catalog(force=True)
        return _json({"status": "ok", "productos": len(_CACHE["catalog"] or [])}, 200)
    except Exception as e:
        return _json({"error": str(e)}, 500)


@app.route('/health', methods=['GET'])
def health():
    # Autodiagnóstico: ¿el cliente service_role puede VER la tabla jobs?
    # Si esto falla o regresa null, la llave SUPABASE_SERVICE_KEY está mal
    # (es de otro proyecto, o es la anon), y por eso no avanza nada.
    svc = {"service_key_set": bool(SUPABASE_SERVICE_KEY)}

    # ¿A qué proyecto de Supabase está conectado Railway? (el host NO es secreto)
    try:
        host = (SUPABASE_URL or "").split("//")[-1].split(".")[0]
        svc["proyecto"] = host
    except Exception:
        svc["proyecto"] = None

    try:
        sb = _new_supabase(service=True)
        rj = sb.table('jobs').select('id', count='exact').limit(1).execute()
        svc["jobs_visibles"] = getattr(rj, 'count', None)
        # Cuenta products con LA MISMA llave service, para confirmar que esta
        # conexión ve el mismo proyecto donde está tu catálogo.
        rp = sb.table('products').select('CodigoArt', count='exact').limit(1).execute()
        svc["products_visibles"] = getattr(rp, 'count', None)
    except Exception as e:
        svc["jobs_error"] = str(e)[:300]

    return _json({
        "status": "ok",
        "catalogo_cargado": _CACHE["catalog"] is not None,
        "productos": len(_CACHE["catalog"] or []),
        "extractor": "on",
        "docling_url": DOCLING_OCR_URL,
        "llm": bool(ANTHROPIC_API_KEY),
        "supabase_service": svc,
    }, 200)


# Precalentar la caché del catálogo al arrancar.
try:
    load_catalog()
except Exception as _e:
    print(f"[startup] no se pudo precargar el catálogo: {_e}")


if __name__ == '__main__':
    port = int(os.environ.get("PORT", 8080))
    app.run(host='0.0.0.0', port=port)
